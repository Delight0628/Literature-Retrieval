"""搜索相关 API 路由

本地知识库优先，无缓存时触发爬虫
"""

import sys
import os

# 确保 backend 目录在 path 中
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from knowledge.search_engine import search_works, get_work_modules, get_module_detail
from knowledge.local_store import init_store, work_exists
from crawler.scraper import MultiSourceCrawler, LITERARY_MODULES

# 初始化知识库
init_store()

router = APIRouter(prefix="/api", tags=["search"])

# 全局爬虫实例
_crawler = None


def _get_crawler() -> MultiSourceCrawler:
    """获取全局爬虫实例"""
    global _crawler
    if _crawler is None:
        _crawler = MultiSourceCrawler()
    return _crawler


class SearchRequest(BaseModel):
    query: str


class DeepSearchRequest(BaseModel):
    query: str
    module_id: str


class DownloadRequest(BaseModel):
    query: str
    module_id: str
    content: str


@router.post("/search")
async def search(request: SearchRequest):
    """泛化检索 - 返回模块列表和概要

    流程：
    1. 本地知识库搜索
    2. 无结果时触发爬虫
    3. 爬取结果自动保存到本地
    """
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="搜索关键词不能为空")

    query = request.query.strip()

    # 1. 先从本地知识库搜索
    results = search_works(query)

    if results:
        # 找到匹配的作品，返回模块列表
        work = results[0]
        modules_result = get_work_modules(work["id"])
        return modules_result

    # 2. 本地无缓存，触发爬虫
    try:
        crawler = _get_crawler()
        crawler_result = crawler.search(query)

        if crawler_result and crawler_result.get("modules"):
            return {
                "query": query,
                "modules": crawler_result["modules"],
                "title": crawler_result.get("title", query),
                "work_url": crawler_result.get("work_url", ""),
                "message": "从互联网获取",
            }
    except Exception as e:
        print(f"爬取失败: {e}")

    # 3. 爬取也失败，返回空结果
    return {
        "query": query,
        "modules": [
            {
                "id": m["id"],
                "name": m["name"],
                "summary": "暂无概要信息",
                "source": "系统默认",
                "source_url": "",
            }
            for m in LITERARY_MODULES
        ],
        "message": f"未找到与「{query}」相关的文学作品",
    }


@router.post("/search/deep")
async def search_deep(request: DeepSearchRequest):
    """深度检索 - 获取指定模块的详细内容

    流程：
    1. 本地知识库查询
    2. 无详情时触发爬虫
    3. 爬取结果自动保存到本地
    """
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="搜索关键词不能为空")

    query = request.query.strip()
    module_id = request.module_id

    # 1. 先从本地知识库查询
    results = search_works(query)
    if results:
        work = results[0]
        detail = get_module_detail(work["id"], module_id)
        if "content" in detail and detail["content"]:
            return detail

    # 2. 本地无详情，触发爬虫
    try:
        crawler = _get_crawler()
        detail = crawler.deep_search(query, module_id)

        if detail and detail.get("content"):
            return detail
    except Exception as e:
        print(f"爬取失败: {e}")

    # 3. 爬取也失败，返回默认提示
    module_info = next(
        (m for m in LITERARY_MODULES if m["id"] == module_id),
        {"id": module_id, "name": module_id, "keywords": []}
    )

    return {
        "module": module_info,
        "content": f"暂未找到关于「{query}」的{module_info['name']}详细信息。\n\n建议尝试其他关键词或稍后重试。",
        "sources": [],
        "images": [],
    }


def _get_module_name(module_id: str) -> str:
    """获取模块中文名称"""
    for m in LITERARY_MODULES:
        if m["id"] == module_id:
            return m["name"]
    return module_id


@router.post("/download")
async def download(request: DownloadRequest):
    """生成文档下载"""
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    import urllib.parse

    doc = Document()

    module_name = _get_module_name(request.module_id)
    title = doc.add_heading(f"{request.query} - {module_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(request.content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 中文文件名需要 URL 编码
    filename = f"{request.query}_{module_name}.docx"
    encoded_filename = urllib.parse.quote(filename)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        },
    )


@router.get("/health")
async def health():
    """健康检查"""
    return {"status": "ok"}
